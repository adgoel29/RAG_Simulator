# def fileload(path):
#     with open(path,'r',encoding="utf-8") as f:
#         con=f.read()
#     return con

import os
import pandas as pd
from pypdf import PdfReader
from docx import Document
import pdfplumber

SUPPORTED = {".pdf", ".txt", ".docx", ".xlsx", ".csv"}


# ─────────────────────────────────────────
# Shared utilities
# ─────────────────────────────────────────

def chunk_text(text, chunk_size=200, overlap=50):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = words[i:i + chunk_size]
        if len(chunk) < 10:
            continue
        chunks.append(" ".join(chunk))
    return chunks


def serialize_table(df, source_name=""):
    """Convert a dataframe to a descriptive text chunk for embedding."""
    if df is None or df.empty:
        return None

    parts = [f"Table from: {source_name}"]
    cols = [str(c) for c in df.columns.tolist()]
    parts.append(f"Columns: {', '.join(cols)}")

    sample = df.head(3).to_string(index=False)
    parts.append(f"Sample rows:\n{sample}")

    numeric_cols = df.select_dtypes(include='number').columns.tolist()
    if numeric_cols:
        stats = df[numeric_cols].describe().loc[['mean', 'min', 'max']]
        parts.append(f"Numeric summary:\n{stats.to_string()}")

    cat_cols = df.select_dtypes(include='object').columns.tolist()
    for col in cat_cols[:4]:
        unique_vals = df[col].dropna().unique()[:8]
        parts.append(f"'{col}' contains: {', '.join(str(v) for v in unique_vals)}")

    return "\n".join(parts)


def _filename_fallback(file_path):
    name = os.path.splitext(os.path.basename(file_path))[0]
    return [name.replace("_", " ").replace("-", " ")]


# ─────────────────────────────────────────
# Per-filetype readers — all return [chunks]
# ─────────────────────────────────────────

def _read_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return chunk_text(text) or _filename_fallback(file_path)
    except Exception as e:
        print(f"WARNING: could not read TXT {file_path}: {e}")
        return _filename_fallback(file_path)


def _read_docx(file_path):
    fname = os.path.basename(file_path)
    try:
        doc = Document(file_path)
        chunks = []

        # Text paragraphs
        full_text = " ".join(
            p.text.strip() for p in doc.paragraphs if p.text.strip()
        )
        if full_text:
            chunks.extend(chunk_text(full_text))

        # Tables
        for table in doc.tables:
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows
            ]
            if not rows:
                continue
            if len(rows) > 1:
                df = pd.DataFrame(rows[1:], columns=rows[0])
            else:
                df = pd.DataFrame(rows)

            serialized = serialize_table(df, source_name=fname)
            if serialized:
                chunks.append(serialized)

        return chunks or _filename_fallback(file_path)

    except Exception as e:
        print(f"WARNING: could not read DOCX {file_path}: {e}")
        return _filename_fallback(file_path)


def _read_pdf(file_path):
    fname = os.path.basename(file_path)
    chunks = []

    try:
        # pdfplumber handles both text and tables
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:

                # Tables first
                for table in page.extract_tables():
                    if not table:
                        continue
                    if len(table) > 1:
                        df = pd.DataFrame(table[1:], columns=table[0])
                    else:
                        df = pd.DataFrame(table)
                    serialized = serialize_table(df, source_name=fname)
                    if serialized:
                        chunks.append(serialized)

                # Text
                text = page.extract_text()
                if text and text.strip():
                    chunks.extend(chunk_text(text))

    except Exception as e:
        print(f"WARNING: pdfplumber failed for {file_path}: {e}, trying pypdf fallback")
        # Fallback to pypdf for text-only
        try:
            reader = PdfReader(file_path)
            text = "".join(
                page.extract_text() or "" for page in reader.pages
            )
            if text.strip():
                chunks.extend(chunk_text(text))
        except Exception as e2:
            print(f"WARNING: pypdf also failed for {file_path}: {e2}")

    return chunks or _filename_fallback(file_path)


def _read_excel(file_path):
    fname = os.path.basename(file_path)
    chunks = []
    try:
        xl = pd.ExcelFile(file_path)
        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            # Use sheet name as source identifier
            source = f"{fname}::{sheet}" if len(xl.sheet_names) > 1 else fname
            serialized = serialize_table(df, source_name=source)
            if serialized:
                chunks.append(serialized)
    except Exception as e:
        print(f"WARNING: could not read Excel {file_path}: {e}")

    return chunks or _filename_fallback(file_path)


def _read_csv(file_path):
    fname = os.path.basename(file_path)
    try:
        df = pd.read_csv(file_path)
        serialized = serialize_table(df, source_name=fname)
        return [serialized] if serialized else _filename_fallback(file_path)
    except Exception as e:
        print(f"WARNING: could not read CSV {file_path}: {e}")
        return _filename_fallback(file_path)


# ─────────────────────────────────────────
# Main entry point
# ─────────────────────────────────────────

_READERS = {
    ".txt":  _read_txt,
    ".docx": _read_docx,
    ".pdf":  _read_pdf,
    ".xlsx": _read_excel,
    ".xls":  _read_excel,
    ".csv":  _read_csv,
}

def get_content(file_path):
    """
    Returns {filename: [chunk1, chunk2, ...]} for a single supported file.
    """
    result = {}

    ext = os.path.splitext(file_path)[1].lower()

    if ext not in _READERS:
        return result

    chunks = _READERS[ext](file_path)

    if chunks:
        filename = os.path.basename(file_path)
        result[filename] = chunks

    return result[filename][0] if result[filename] else ""


if __name__ == "__main__":
    ans = get_content("/home/aditya/Desktop/RAG_Simulator/chemistry project.docx")
    print(ans)